#!/usr/bin/env python3
"""Build docs/Work_Management_Manual.docx — the user-facing operations manual.

Regenerate after feature changes:  python3 scripts/build_manual.py
The content lives here as structured data so the manual grows with the system.
"""
import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = pathlib.Path(__file__).parent.parent
OUT = ROOT / "docs" / "Work_Management_Manual.docx"

INK = RGBColor(0x0A, 0x0A, 0x0A)
GOLD = RGBColor(0xA0, 0x60, 0x00)
MUTE = RGBColor(0x6F, 0x6D, 0x65)

# ── manual content: (kind, text) — kind: h1/h2/h3/p/b (bullet)/n (numbered) ──
SECTIONS = [
    ("h1", "Upande Work Management — Operations Manual"),
    ("p", "This manual explains how to plan, assign, record, audit and pay farm task work "
          "using the Work Management system, and how the system protects the payroll with "
          "attendance checks and standing audits. Pages live on your site: /work-management "
          "(dashboard), /work-planner, /work-assigner, /work-actuals, /work-payment."),
    ("p", "Sign in with your normal account. The round button at the top-right of every page "
          "shows who is signed in and is used to log in and out."),

    ("h2", "1. The pipeline at a glance"),
    ("p", "Every shilling follows the same path: Plan → Assign → Actuals → Confirm → Pay. "
          "A plan sets the task, blocks, period, daily standard and rate. An assignment puts "
          "named workers on an approved plan. Actuals record each worker's daily output. "
          "Confirmations (Farm Manager → HR → General Manager) turn records into payable "
          "earnings, and Payment reviews and releases the money one worker at a time."),

    ("h2", "2. Planner (/work-planner)"),
    ("n", "Pick the farm, task, block or blocks, the work period, and the crew size."),
    ("n", "The daily standard and rate come from the task list (for example 150 Meter/day at "
          "KES 2.2667 per Meter — rates carry four decimals so a full day comes to exactly KES 340)."),
    ("n", "Save. The plan starts as PENDING and goes to the farm's approver (Farm Manager)."),
    ("n", "Approved plans become available to the Assigner. A plan's budget = rate × target quantity."),

    ("h2", "3. Assigner (/work-assigner)"),
    ("n", "Pick an approved plan; the farm's workers load in the picker."),
    ("n", "Every worker carries live badges: presence (P · 06:12 scanned in, 'not in yet', or "
          "'night shift'), off days in the window, leave, absences, and 'assigned elsewhere' "
          "(a worker on another live assignment for an overlapping period cannot be picked — "
          "this prevents double allocation and double pay)."),
    ("n", "When the window includes today, a presence bar shows how many of the farm's workers "
          "have scanned in, with a 'only workers who are in' filter and a Refresh scans link."),
    ("n", "Selecting a flagged worker asks for explicit confirmation; submitting re-checks on "
          "the server and logs every override on the assignment."),
    ("n", "The General Manager signs off assignments. Mid-job changes use the swap button — "
          "substitutions record who left, who joined and when."),

    ("h2", "4. Actuals (/work-actuals)"),
    ("n", "Open the assignment; the grid shows one row per worker and one column per day."),
    ("n", "Every past/today cell carries presence evidence: the check-in time (in 06:52), "
          "P (marked present, no scan time), A · absent (marked Absent), or ? (no record either "
          "way — presence unknown). Rest days show a dot, approved leave days are blocked."),
    ("n", "Enter each worker's daily quantity. Rows are valued at qty × rate to the cent; on "
          "plans whose implied daily wage is within 1% of KES 340 a full day is valued at "
          "exactly 340.00."),
    ("n", "Saving warns when a quantity conflicts with attendance: recording work for a worker "
          "marked Absent needs the Farm Manager (or GM) to approve; leave/off/no-scan conflicts "
          "can be overridden by the enterer, and every override is logged on the document. "
          "Saving also warns when the same worker-day is already recorded for the task in "
          "another document (double pay)."),
    ("n", "Submit walks the document through Farm Manager → HR → GM to CONFIRMED. Submission "
          "unlocks only when the plan target is reached; plans that cannot finish (absentees, "
          "crop finished early) are closed early via a close request, which the GM approves — "
          "the close queue shows live done/remaining figures."),

    ("h2", "5. Payment (/work-payment) — workers are paid one at a time"),
    ("p", "The payment section is worker-centric. Each worker's confirmed earnings are sent "
          "to accounts as their own payment entry and released. The status ladder per worker "
          "is: Unpaid → Sent to accounts → Paid. There is no separate review step — sending "
          "IS the sign-off: the sender's name and time are stamped on every included day-row "
          "and on the payment entry. The review sheet stays available for checking anyone "
          "before sending."),
    ("h3", "5.1 Reviewing a worker"),
    ("b", "Click any worker to open the review sheet: identity, window KPIs (earned, paid, "
          "unpaid, days), one card per task with the daily log underneath, a Payments tab, and "
          "a Discrepancies tab listing that worker's flagged days."),
    ("b", "Each task card names the full accountability chain: who created the plan, assigned "
          "the job, captured the actuals, and each approver (FM, HR, GM), plus the task's "
          "standard (e.g. 300 Tree/day @ KES 1.1333)."),
    ("b", "Every day row shows presence evidence next to the pay. In Work & days, every "
          "unpaid day's quantity is directly editable — change as many as needed and press "
          "the single Save changes button at the bottom (Undo restores the originals). Pay "
          "recomputes at each row's rate, documents re-sum, and one audit comment per "
          "document lists every change."),
    ("b", "Download Excel exports the review as a workbook: a Summary sheet and a Tasks & days "
          "sheet laid out like the review (one table per task with its day rows and presence)."),
    ("h3", "5.2 Sending to accounts"),
    ("b", "Submit & send to accounts creates one payment document for that worker (WMPAY-…) "
          "holding the period, totals, who sent it and when, and one line per actuals document "
          "with the task, block, worked period, days, qty, rate, amount and the whole sign-off "
          "chain. The sender is stamped as reviewer on every included day-row."),
    ("b", "Bulk: tick several workers (workers with attendance conflicts carry a red flag with "
          "the day count) and use Send to accounts — each still gets their own entry."),
    ("h3", "5.3 Awaiting accounts"),
    ("b", "Accounts releases an entry with Mark paid — every included day row is stamped paid."),
    ("b", "Return to unpaid withdraws an entry (deletes the reference, clears review stamps) so "
          "the days can be corrected and re-sent. Bulk return handles many entries at once."),

    ("h2", "6. Time & attendance protection"),
    ("p", "The pipeline checks workers against attendance before work is given to them or "
          "recorded for them. Every behaviour is a switch in Work Management Settings."),
    ("b", "Checks: marked Absent (submitted attendance), approved leave, weekly offs/holidays, "
          "and morning presence (biometric scan or Present attendance today, with a "
          "configurable cutoff so early assigning is never blocked; night shifts are exempt)."),
    ("b", "The off-day rule when assigning is configurable: flag only when offs cover the whole "
          "window (default), flag any off day, or ignore offs at assignment."),
    ("b", "Overrides: leave/off/no-scan conflicts can be overridden by the person doing the "
          "work, always logged. Absent-day actuals need the Farm Manager or GM."),
    ("b", "Missing attendance never blocks anyone — only an explicit Absent record does, so "
          "device sync gaps cannot stop work."),

    ("h2", "7. Discrepancies — the standing audit"),
    ("p", "Payment → Audit → Discrepancies scans every confirmed worker-day in the chosen "
          "window and groups everything suspicious. Each check has its own settings checkbox; "
          "rows link to the worker's review sheet and name the documents."),
    ("b", "Paid on marked-Absent days — a scan time means the attendance record is probably "
          "wrong; no scan means the entry needs scrutiny. Days whose attendance was corrected "
          "(a Present record exists alongside an old Absent one) are validated out and never "
          "flagged."),
    ("b", "No presence evidence at all — no scan and no attendance record of any kind."),
    ("b", "Earning while on approved leave — possible double payment."),
    ("b", "Work on off days / holidays — fine if deliberate overtime."),
    ("b", "Amount ≠ qty × rate — edited or corrupted values."),
    ("b", "Two farms, one day — physically doubtful."),
    ("b", "Entered and approved by the same person — no independent check."),
    ("b", "Earning after leaving the job — days dated after a worker was released."),
    ("b", "Paid twice for the same day — the same worker, task and date in two documents; a "
          "one-click repair keeps the earliest copy and zeroes the rest."),
    ("b", "Recorded work with no pay — rows valued at zero although the document has a rate; a "
          "one-click Revalue repairs them at qty × rate."),
    ("b", "Left the company but still on live assignments — ex-employees still assignable. A "
          "Clean slate button releases the backlog; the auto-release setting (off by default) "
          "releases future leavers the moment HR deactivates them. Released workers keep all "
          "recorded work and pay — release only stops new quantities."),

    ("h2", "8. Dashboard (/work-management)"),
    ("b", "Activity across the pipeline — stage cards for Planned, Assigned (with active "
          "employees split into task workers and permanent staff), Actual and Payment."),
    ("b", "Workers & value per farm — per-farm cards: assigned workers, active employees "
          "(task/permanent split), awaiting actuals, confirmed, quantities and value."),
    ("b", "Delivery timeline — planned vs staffed vs delivered per day, with farm and date "
          "filters and quantity/KES toggle."),
    ("b", "Field intelligence (beside the timeline) — two tabs: Efficiency (Ha per man-day "
          "and Cost per Ha, per farm and per task) and Available workers (the number and list "
          "of employees free to work on any chosen date, filterable by farm). Explained in "
          "full in the next section."),
    ("b", "Action queues — everything waiting on someone, one queue at a time."),
    ("b", "Approver KPIs — per approval stage: each approver's sign-offs, value and time taken."),
    ("b", "Value flow — weekly planned/assigned/confirmed value and a per-plan table with each "
          "plan's accountability chain."),
    ("b", "Pipeline performers — planner and assigner economics, per person, with "
          "most/least-expensive callouts (judged only on people with real volume, more than "
          "500 units). Column key: Plans = approved plans created in the window; Target qty = "
          "the output those plans promised; Actual qty = confirmed output delivered; "
          "Achieved = Actual ÷ Target (green 90%+, amber 60%+, red below); Budget KES = what "
          "the plans are worth if fully delivered (rate × target); Spent KES = confirmed pay "
          "earned on them; Of budget = Spent ÷ Budget — low is NOT automatically savings, "
          "read it with Achieved (50% spent at 50% achieved just means half the work "
          "happened); KES/unit = Spent ÷ Actual, what one unit of output cost under this "
          "person. The Assigners tab uses the same definitions over their assignments, plus "
          "Workers put on jobs (assignment rows they created — a worker on two assignments "
          "counts twice)."),
    ("b", "Crew movements — substitution history: who left, who joined, swaps."),

    ("h2", "9. Field intelligence, explained"),
    ("p", "The Field intelligence card sits beside the Delivery timeline on the dashboard. It "
          "answers two everyday management questions: 'what does our work actually cost per "
          "hectare, and how much ground does a worker-day cover?' and 'who is free to work "
          "today (or any day I pick)?'"),
    ("h3", "9.1 Efficiency tab"),
    ("p", "Pick a date window (default: the last 30 days) and press Apply. The table shows, "
          "per farm and in total:"),
    ("b", "Area Ha — the hectares of the blocks whose plans had confirmed work in the window. "
          "Each plan's blocks are counted once for that plan, using the Area (HA) captured on "
          "the block record (Warehouse). If a block has no area captured, it contributes "
          "nothing — a dash (—) in the table means areas are missing, and the fix is data "
          "entry on the block records, not a system fault."),
    ("b", "Man-days — one worker working one day is one man-day, counted from confirmed "
          "actuals (a worker on two tasks the same day is still one man-day)."),
    ("b", "Ha / man-day — Area ÷ man-days: how much ground one worker-day covers. Higher is "
          "leaner. Compare farms with care: the task mix matters (a farm doing slow detailed "
          "tasks like handling will always cover fewer hectares per man-day than one doing "
          "slashing)."),
    ("b", "Cost KES — the confirmed pay for that work in the window."),
    ("b", "Cost / Ha — Cost ÷ Area: what a hectare of work cost. This is the number to watch "
          "over time per farm and per task. When Cost/Ha rises on the same task mix, each "
          "hectare is consuming more paid work than before. Read it together with Ha/man-day: "
          "if Cost/Ha rises while Ha/man-day falls, workers are covering less ground per day — "
          "productivity dropped (denser weeds, harder terrain, crop stage slowing the task, or "
          "quantities recorded that don't match ground actually covered). If Ha/man-day is "
          "steady while Cost/Ha rises, the change is in the rates — check the task list."),
    ("p", "Below the farm table, the same metrics appear per task (top tasks by man-days), "
          "which is where differences usually explain themselves — compare the same task "
          "across time, not different tasks against each other."),
    ("p", "Worked example: Endebess confirmed 7,871 man-days over blocks totalling 1,670 Ha "
          "at a cost of KES 2.40M → 0.212 Ha per man-day and KES 1,439 per Ha."),
    ("h3", "9.2 Available workers tab"),
    ("p", "Pick any date (past, today or future) and optionally a farm. The card shows the "
          "count and the full name list of AVAILABLE workers — active employees with no live "
          "assignment covering that date (assignment start and leaving dates are respected). "
          "For today and past dates, a green P marks workers who scanned in that day, so you "
          "can see at a glance who is both free AND on the farm. Use the name search to find "
          "someone specific."),
    ("p", "This is the assigner's shortlist: when a new plan needs a crew, the Available "
          "workers list for the start date is exactly who can be picked without a "
          "double-allocation conflict."),

    ("h2", "10. Work Management Settings — reference"),
    ("b", "Attendance gates: check attendance (Absent), check approved leaves, check weekly "
          "offs/holidays, off-day rule when assigning."),
    ("b", "Morning presence: require a morning scan for day-of assignment, morning scan cutoff "
          "(default 09:00), check scans when recording actuals."),
    ("b", "Auto-release inactive employees from assignments (off by default)."),
    ("b", "Discrepancy checks: one checkbox per check listed in section 7."),

    ("h2", "11. Definitions"),
    ("b", "Man-day — one worker working one day (a worker on two tasks the same day is one man-day)."),
    ("b", "Ha / man-day — area of the blocks whose plans were worked, divided by the man-days "
          "spent on them: how much ground one worker-day covers."),
    ("b", "Cost / Ha — confirmed pay divided by the area worked: what a hectare of work costs."),
    ("b", "Available worker — an active employee with no live assignment covering the chosen date."),
    ("b", "Standard — the plan's daily expectation, e.g. 150 Meter/day @ KES 2.2667."),
    ("b", "Presence evidence — a biometric check-in time, a Present attendance record, an "
          "Absent record, or nothing (? — unknown)."),
]


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    for kind, text in SECTIONS:
        if kind == "h1":
            p = doc.add_heading(text, level=0)
            for r in p.runs:
                r.font.color.rgb = INK
            sub = doc.add_paragraph("Kaitet Group · kaitet-group.upande.com · living document — regenerated with the system")
            sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in sub.runs:
                r.font.color.rgb = MUTE
                r.font.size = Pt(9)
        elif kind == "h2":
            p = doc.add_heading(text, level=1)
            for r in p.runs:
                r.font.color.rgb = GOLD
        elif kind == "h3":
            p = doc.add_heading(text, level=2)
            for r in p.runs:
                r.font.color.rgb = INK
        elif kind == "b":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "n":
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(text)
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    build()
